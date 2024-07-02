import sys
import datetime
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, QTextEdit, QFileDialog, QMessageBox, QComboBox
from PyQt5.QtGui import QFont  # Import QFont to adjust font sizes
from docx import Document

class ShopApplication(QWidget):
    def __init__(self):
        super().__init__()
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("Bhutta Sanitary Store")
        self.setGeometry(100, 100, 1000, 800)  # Increased window size

        # Create and set font size for labels
        label_font = QFont()
        label_font.setPointSize(14)

        # Create and set font size for entries and buttons
        entry_font = QFont()
        entry_font.setPointSize(12)

        button_font = QFont()
        button_font.setPointSize(12)

        # Create labels and set fonts
        self.customer_name_label = QLabel("Customer Name:")
        self.customer_name_label.setFont(label_font)
        self.customer_name_entry = QLineEdit()
        self.customer_name_entry.setFont(entry_font)
        self.customer_name_entry.setPlaceholderText("Enter customer name")

        self.category_label = QLabel("Category:")
        self.category_label.setFont(label_font)
        self.category_combo = QComboBox()
        self.category_combo.setFont(entry_font)
        self.category_combo.addItem("Select Category")
        self.category_combo.addItem("Popular Pn 16")
        self.category_combo.addItem("Sewerage")
        self.category_combo.currentIndexChanged.connect(self.load_subcategories)

        self.subcategory_label = QLabel("Subcategory:")
        self.subcategory_label.setFont(label_font)
        self.subcategory_combo = QComboBox()
        self.subcategory_combo.setFont(entry_font)
        self.subcategory_combo.addItem("Select Subcategory")
        self.subcategory_combo.currentIndexChanged.connect(self.load_items)

        self.item_label = QLabel("Item:")
        self.item_label.setFont(label_font)
        self.item_combo = QComboBox()
        self.item_combo.setFont(entry_font)
        self.item_combo.addItem("Select Item")

        self.item_quantity_label = QLabel("Quantity:")
        self.item_quantity_label.setFont(label_font)
        self.item_quantity_entry = QLineEdit()
        self.item_quantity_entry.setFont(entry_font)
        self.item_quantity_entry.setPlaceholderText("Quantity")

        self.add_item_button = QPushButton("Add Item")
        self.add_item_button.setFont(button_font)
        self.add_item_button.clicked.connect(self.add_item)

        self.items_list = QTextEdit()
        self.items_list.setFont(entry_font)
        self.items_list.setReadOnly(True)
        self.items_list.setFixedHeight(200)  # Set a fixed height for the items list

        self.discount_label = QLabel("Discount (%):")
        self.discount_label.setFont(label_font)
        self.discount_entry = QLineEdit()
        self.discount_entry.setFont(entry_font)
        self.apply_discount_button = QPushButton("Apply Discount")
        self.apply_discount_button.setFont(button_font)
        self.apply_discount_button.clicked.connect(self.apply_discount)

        self.generate_bill_button = QPushButton("Generate Bill")
        self.generate_bill_button.setFont(button_font)
        self.generate_bill_button.clicked.connect(self.generate_bill)

        self.invoice_text = QTextEdit()
        self.invoice_text.setFont(entry_font)
        self.invoice_text.setReadOnly(True)
        self.invoice_text.setFixedHeight(200)  # Set a fixed height for the invoice text

        # Create layout and add widgets
        self.layout = QVBoxLayout()
        self.layout.addWidget(self.customer_name_label)
        self.layout.addWidget(self.customer_name_entry)
        self.layout.addWidget(self.category_label)
        self.layout.addWidget(self.category_combo)
        self.layout.addWidget(self.subcategory_label)
        self.layout.addWidget(self.subcategory_combo)
        self.layout.addWidget(self.item_label)
        self.layout.addWidget(self.item_combo)
        self.layout.addWidget(self.item_quantity_label)
        self.layout.addWidget(self.item_quantity_entry)
        self.layout.addWidget(self.add_item_button)
        self.layout.addWidget(self.items_list)
        self.layout.addWidget(self.discount_label)
        self.layout.addWidget(self.discount_entry)
        self.layout.addWidget(self.apply_discount_button)
        self.layout.addWidget(self.generate_bill_button)
        self.layout.addWidget(self.invoice_text)

        self.setLayout(self.layout)

        self.items = []

    def load_subcategories(self):
        selected_category = self.category_combo.currentText()
        self.subcategory_combo.clear()
        self.subcategory_combo.addItem("Select Subcategory")

        if selected_category != "Select Category":
            subcategories = self.category_data[selected_category]["Subcategories"]
            self.subcategory_combo.addItems(subcategories.keys())

    def load_items(self):
        selected_category = self.category_combo.currentText()
        selected_subcategory = self.subcategory_combo.currentText()
        self.item_combo.clear()
        self.item_combo.addItem("Select Item")

        if (
            selected_category != "Select Category"
            and selected_subcategory != "Select Subcategory"
        ):
            items = self.category_data[selected_category]["Subcategories"].get(
                selected_subcategory, []
            )
            self.item_combo.addItems(items)

    def add_item(self):
        item_name_with_price = self.item_combo.currentText()
        item_quantity = self.item_quantity_entry.text()

        if item_name_with_price == "Select Item":
            QMessageBox.warning(self, "Error", "Please select an item.")
            return

        if not item_quantity:
            QMessageBox.warning(self, "Error", "Please enter item quantity.")
            return

        item_name, item_price = self.extract_item_info(item_name_with_price)
        item = {"name": item_name, "price": item_price, "quantity": int(item_quantity)}
        self.items.append(item)
        self.update_items_list()
        self.clear_input_fields()

    def update_items_list(self):
        self.items_list.clear()
        for item in self.items:
            self.items_list.append(f"{item['name']} - {item['quantity']} x {item['price']:.2f} Rupees")

    def clear_input_fields(self):
        self.item_combo.setCurrentIndex(0)
        self.item_quantity_entry.clear()

    def generate_bill(self):
        customer_name = self.customer_name_entry.text()
        if not customer_name:
            QMessageBox.warning(self, "Error", "Please enter customer name.")
            return

        if not self.items:
            QMessageBox.warning(self, "Error", "No items in the cart.")
            return

        total_cost = sum(item["price"] * item["quantity"] for item in self.items)
        discount = float(self.discount_entry.text() or 0)
        discounted_cost = total_cost - (total_cost * discount / 100)

        document = Document()
        document.add_heading(f"Bill for {customer_name}", 0)
        document.add_heading(datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), level=1)
        document.add_paragraph("\nItems:")

        for item in self.items:
            document.add_paragraph(f"{item['name']} - {item['quantity']} x {item['price']:.2f} Rupees")

        document.add_paragraph(f"\nTotal Cost: {total_cost:.2f} Rupees")
        document.add_paragraph(f"Discount: {discount}%")
        document.add_paragraph(f"Discounted Cost: {discounted_cost:.2f} Rupees")

        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getSaveFileName(self, "Save Bill", "", "Word Files (*.docx)")

        if file_name:
            document.save(file_name)
            QMessageBox.information(self, "Bill Generated", f"Bill saved as {file_name}")

        # Generate and display the invoice
        self.generate_invoice(total_cost, discount, discounted_cost)

    def apply_discount(self):
        pass  # Implement discount functionality here

    def extract_item_info(self, item_name_with_price):
        parts = item_name_with_price.split(" - ")
        item_name = parts[0]
        item_price = float(parts[1].replace(" ", ""))
        return item_name, item_price

    def generate_invoice(self, total_cost, discount, discounted_cost):
        invoice_text = f"Customer Name: {self.customer_name_entry.text()}\n"
        invoice_text += f"Date: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        invoice_text += "Items:\n"

        for item in self.items:
            invoice_text += f"{item['name']} - {item['quantity']} x {item['price']:.2f}\n"

        invoice_text += f"\nTotal Cost: {total_cost:.2f} Rupees\n"
        invoice_text += f"Discount: {discount}%\n"
        invoice_text += f"Discounted Cost: {discounted_cost:.2f} Rupees"

        self.invoice_text.setPlainText(invoice_text)

def main():
    app = QApplication(sys.argv)
    window = ShopApplication()
    window.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
